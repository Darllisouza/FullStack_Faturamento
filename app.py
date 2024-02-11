from flask import Flask, render_template, request, send_file
import os
import pandas as pd
import pdfplumber
from pathlib import Path
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from docx import Document
from pdfminer.high_level import extract_text
from werkzeug.utils import secure_filename
import time
import locale


app = Flask(__name__)

#Essa parte do código é fundamental para garantir a segurança e o controle dos arquivos que são enviados para o Flask. Limitar as extensões de arquivo permitidas ajuda a prevenir potenciais vulnerabilidades, como a execução de código malicioso ou a sobrecarga do servidor com arquivos indesejados.
#Configurações para uploads
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf'}
# Caminho para o template do Word
TEMPLATE_PATH = './templates/template.docx'


app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/processar_pdf', methods=['POST'])
def processar_pdf():

    file = request.files['file']

        # Verifica se o arquivo foi enviado na requisição
    if 'file' not in request.files:
        return 'Nenhum arquivo enviado'

    # Verifica se o arquivo possui um nome válido
    if file.filename == '':
        return 'Nenhum arquivo selecionado'

    # Verifica se a extensão do arquivo é permitida
    if file and allowed_file(file.filename):
        # Salva o arquivo no diretório de uploads
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        print("Arquivo PDF salvo com sucesso:", file_path)
    else:
        return 'Extensão de arquivo não permitida'

    # Abre o arquivo PDF
    with pdfplumber.open(file_path) as pdf:
        # Variável para armazenar as linhas relevantes
        relevant_rows = []
        # Variável para rastrear se estamos em uma seção relevante
        in_section = False
        # Itera sobre todas as páginas do PDF
        for page in pdf.pages:
            # Extrai todas as tabelas da página
            tables = page.extract_tables()
            # Processa as tabelas extraídas
            for table in tables:
                # Itera sobre as linhas da tabela
                for row in table:
                    # Verifica se estamos em uma seção relevante
                    if row[0] == '2.2.1) Mercado Interno' or row[0] == '2.2.2) Mercado Externo':
                        in_section = True
                        relevant_rows.append(row)
                    # Se estamos em uma seção relevante, adiciona a linha
                    elif in_section:
                        # Verifica se devemos parar de coletar linhas
                        if row[0] == 'Receita Bruta Auferida (regime competência)' or row[0] == 'Valor Total do Débito Declarado (R$)':
                            in_section = False
                            break
                        relevant_rows.append(row)

    # Cria o DataFrame
    df = pd.DataFrame(relevant_rows)


    # Caminho completo para o arquivo Excel
    excel_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'Declaracao_de_faturamento.xlsx')
    df.to_excel(excel_file_path, index=False)

    # Remove ou adiciona colunas necessárias
    indices_remover = []
    for i, row in df.iterrows():
        if "Mercado Interno" in str(row.values[0]) or "Mercado Externo" in str(row.values[0]):
            indices_remover.append(i)

    df_sem_secoes = df.drop(indices_remover)

    # Define quais colunas devem ser interpretadas como datas
    colunas_datas = [0, 2, 4, 6]  # As colunas A, C, E, G têm índices 0, 2, 4, 6 respectivamente

    # Converte as colunas para datas, tratando valores inválidos como NaT
    for coluna in colunas_datas:
        df[coluna] = pd.to_datetime(df[coluna], format="%m/%Y", errors='coerce').dt.strftime("%m/%Y")

    # Salva o DataFrame de volta no arquivo Excel
    df.to_excel(excel_file_path, index=False)

    # Define as colunas de data e valor
    colunas_datas = [0, 2, 4, 6]
    colunas_valores = [1, 3, 5, 7]

    # Inicializa listas para armazenar datas e valores de faturamento
    datas = []
    valores = []

    # Itera sobre as linhas do DataFrame
    for index, row in df.iterrows():
        # Itera sobre as colunas de datas e valores
        for col_data, col_valor in zip(colunas_datas, colunas_valores):
            # Verifica se o valor na célula não está vazio
            if pd.notnull(row[col_data]) and pd.notnull(row[col_valor]):
                # Remove pontos de milhar e substitui vírgula por ponto
                valor = row[col_valor].replace(".", "").replace(",", ".")
                # Adiciona a data e o valor convertido para ponto flutuante
                datas.append(pd.to_datetime(row[col_data], format='%m/%Y'))  # Converte para datetime
                valores.append(float(valor.replace(",", ".")))  # Substitui a vírgula por ponto e converte para float

    # Cria um novo DataFrame com as datas e valores de faturamento
    df_final = pd.DataFrame({"Data": datas, "Valor": valores})

    # Ordena o DataFrame pela coluna de datas
    df_order = df_final.sort_values(by='Data')

    # Agrupa os dados pela coluna 'Data' e soma os valores correspondentes
    df_agrupado = df_order.groupby('Data')['Valor'].sum().reset_index()

    # Ordena o DataFrame agrupado pela coluna 'Data'
    df_agrupado = df_agrupado.sort_values(by='Data')

    # Salva o DataFrame de volta no arquivo Excel
    df_agrupado.to_excel(excel_file_path, index=False)

    # Encontra a data mais recente e define a data inicial para 12 meses antes dela
    data_atual = pd.to_datetime('today')
    data_inicial = data_atual - pd.DateOffset(months=13)

    # Filtra o DataFrame agrupado para incluir apenas os últimos 12 meses
    df_ultimos_12_meses = df_agrupado[df_agrupado['Data'] >= data_inicial]

    # Salva o DataFrame de volta no arquivo Excel
    df_ultimos_12_meses.to_excel(excel_file_path, index=False)

    # Achar o valor do mês atual
    excel_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'Declaracao_de_faturamento.xlsx')

    valor_total = None  # Inicializa a variável como None
    try:
        # Abre o arquivo PDF e procura pelo valor total
        with pdfplumber.open(file_path) as pdf:
            # Itera sobre todas as páginas do PDF
            for page in pdf.pages:
                # Extrai todo o texto da página
                text = page.extract_text()

                # Divide o texto em linhas
                lines = text.split('\n')

                # Itera sobre as linhas do texto
                for line in lines:
                    # Verifica se a linha contém a informação desejada
                    if "Receita Bruta do PA (RPA) - Competência" in line:
                        # Divide a linha pelos espaços em branco
                        values = line.split()

                        # Verifica se a linha contém pelo menos 4 valores
                        if len(values) >= 4:
                            # O valor total é o quarto valor na linha
                            valor_total = values[-1]  # Último valor na linha
                            break  # Sai do loop após encontrar o valor total
                if valor_total:
                    break
    except Exception as e:
        print("Erro ao abrir o arquivo PDF:", e)

    # Obtém a última data no DataFrame df_ultimos_12_meses
    try:
        # Carrega o DataFrame df_ultimos_12_meses existente do arquivo Excel
        df_ultimos_12_meses = pd.read_excel(excel_file_path)

        # Verifica se o DataFrame não está vazio e se a coluna 'Data' está presente
        if not df_ultimos_12_meses.empty and 'Data' in df_ultimos_12_meses.columns:
            ultima_data = df_ultimos_12_meses['Data'].iloc[-1]

            # Calcula a próxima data para o próximo mês
            proxima_data = ultima_data + pd.DateOffset(months=1)

            # Adiciona uma nova linha ao DataFrame df_ultimos_12_meses para o próximo mês
            nova_linha = {'Data': proxima_data, 'Valor': valor_total}  
            df_ultimos_12_meses = pd.concat([df_ultimos_12_meses, pd.DataFrame([nova_linha])], ignore_index=True)


            # Salva o DataFrame atualizado no arquivo DECLARACAO_DE_FATURAMENTO.xlsx
            df_ultimos_12_meses.to_excel(excel_file_path, index=False)
        else:
            print("Erro ao carregar ou processar o arquivo Excel:", e)
    except Exception as e:
        print("Erro ao carregar ou processar o arquivo Excel:", e)

    # Converter em valor e data
    df_ultimos_12_meses['Valor'] = df_ultimos_12_meses['Valor'].replace('0,00', '0') 
    df_ultimos_12_meses['Valor'] = df_ultimos_12_meses['Valor'].apply(lambda x: locale.format_string('%.2f', float(x), grouping=True))
    df_ultimos_12_meses['Data'] = df_ultimos_12_meses['Data'].dt.strftime('%m/%Y')
    df_ultimos_12_meses['Valor'] = df_ultimos_12_meses['Valor'].str.replace('.', ',')

    # Salvar planilha declaração de faturamento
    df_ultimos_12_meses.to_excel(excel_file_path, index=False)

    # Abre o arquivo PDF
    with pdfplumber.open(file_path) as pdf:
        # Itera sobre as páginas do PDF
        for page in pdf.pages:
            # Extrai o texto da página
            text = page.extract_text()

            # Verifica se os campos estão presentes no texto extraído
            if "CNPJ Matriz:" in text:
                # Extrai o valor do CNPJ Matriz
                cnpj_matriz = text.split("CNPJ Matriz:")[1].strip().split("\n")[0].strip()

            if "Nome empresarial:" in text:
                # Extrai o valor do Nome Empresarial
                nome_empresarial = text.split("Nome empresarial:")[1].strip().split("\n")[0].strip()

    # Crie um novo DataFrame com as informações da empresa
    df_info_empresa = pd.DataFrame({
        'Nome Empresarial': [nome_empresarial],
        'CNPJ Matriz': [cnpj_matriz]
    })

    # Transpõe o DataFrame de informações da empresa
    df_info_empresa = df_info_empresa.transpose()

    # Caminho para o arquivo Excel
    excel_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'Declaracao_de_faturamento.xlsx')

    with pd.ExcelWriter(excel_file_path) as writer:
        # Escreve as informações da empresa
        df_info_empresa.to_excel(writer, startcol=0, startrow=8, header=False)
        # Escreve os dados de data e valor
        df_ultimos_12_meses.to_excel(writer, startcol=0, startrow=11, index=False, header=["Data", "Valor"])

    return 'Arquivo processado com sucesso'

@app.route('/gerar_documento_word', methods=['POST'])
def gerar_documento_word():
    # Caminho completo para o arquivo Excel
    excel_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'Declaracao_de_faturamento.xlsx')

    if not os.path.exists(excel_file_path):
        print(f"O arquivo Excel '{excel_file_path}' não foi encontrado.")
        exit()

    # Ler o arquivo Excel
    wb = load_workbook(excel_file_path)
    ws = wb.active

    # Mapear as informações relevantes do Excel
    empresa_excel = ws['B9'].value
    cnpj_excel = ws['B10'].value

    # Abrir o arquivo Word
    doc = Document(TEMPLATE_PATH)

    # Procurar pela tabela no arquivo Word e substituir os valores
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Empresa" in cell.text:
                    # Substituir o valor da empresa pelo conteúdo da célula B9 do Excel
                    row.cells[1].text = empresa_excel
                elif "CNPJ" in cell.text:
                    # Substituir o valor do CNPJ pelo conteúdo da célula B10 do Excel
                    row.cells[1].text = cnpj_excel

        # Encontrar a tabela no arquivo Word e substituir os valores
        if len(table.columns) >= 2:
            if table.cell(0, 0).text.strip() == "Período" and table.cell(0, 1).text.strip() == "Faturamento":
                for i, row in enumerate(table.rows[1:], start=13):
                    periodo_excel = ws.cell(row=i, column=1).value
                    faturamento_excel_str = ws.cell(row=i, column=2).value
                    faturamento_excel = float(faturamento_excel_str.replace(",", ".")) if faturamento_excel_str else None
                    if faturamento_excel is not None:
                        if len(f'{faturamento_excel:.2f}'.split('.')[1]) > 2:
                            formatted_faturamento = f"R$ {faturamento_excel:,.2f}".replace(",", ".", 1)
                        else:
                            formatted_faturamento = f"R$ {faturamento_excel:,.2f}".replace(".", ",")
                    row.cells[0].text = str(periodo_excel) if periodo_excel else ""
                    row.cells[1].text = formatted_faturamento if formatted_faturamento else ""
                    for cell_idx, cell in enumerate(row.cells):
                        if cell_idx == 0:
                            cell.paragraphs[0].alignment = 1
                        else:
                            cell.paragraphs[0].alignment = 2

    # Salvar as alterações no arquivo Word
    result_path = os.path.join(app.config['UPLOAD_FOLDER'], 'Demonstrativo de Faturamento.docx')
    doc.save(result_path)

    return 'Código executado com sucesso!'

@app.route('/baixar_documento_word', methods=['GET'])
def baixar_documento_word():
    # Caminho completo para o arquivo Word gerado
    result_path = os.path.join(app.config['UPLOAD_FOLDER'], 'Demonstrativo de Faturamento.docx')

    # Verifica se o arquivo existe
    if os.path.exists(result_path):
        # Retorna o arquivo para download
        return send_file(result_path, as_attachment=True)
    else:
        # Se o arquivo não existe, retorna uma mensagem de erro
        return 'O arquivo não foi encontrado.', 404

if __name__ == '__main__':
    app.run(debug=True)
